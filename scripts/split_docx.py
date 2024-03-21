from docx import Document
from pathlib import Path

# This works: https://products.groupdocs.app/splitter/docx but is hard to automate
# Alternatively, I could try to split the pdf version and convert the pdf pages to docx

def split_docx(docx_path: Path, outpath: Path):
	"""
	Split a docx file into multiple docx files, one for each page.
	"""
	doc = Document(docx_path)
	page = Document()
	j = 0
	done = False
	for paragraph in doc.paragraphs:
		if done:
			page = Document()
			done = False
		# Does this only add the text contents of the paragraph? I want everything
		page.add_paragraph(paragraph.text)
		if paragraph.contains_page_break:
			j += 1
			page.save(outpath / f"{j}.docx")
			done = True

if __name__ == "__main__":
	docx_path = Path("../documents/docxs/test.docx")
	outpath = Path(".")
	split_docx(docx_path, outpath)

